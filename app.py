from flask import Flask, render_template, request, redirect, url_for, flash,send_file,jsonify, session
import json
import mysql.connector
import secrets
from mysql.connector import Error
from docx import Document
from io import BytesIO
from flask import session
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import random
from flask_mail import Mail, Message
import os
from werkzeug.utils import secure_filename 
from werkzeug.security import check_password_hash,generate_password_hash
import jwt
import datetime
import re



app = Flask(__name__)
app.secret_key = secrets.token_hex(16)

UPLOAD_FOLDER = 'static/uploads'  # Folder to store images
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
            
app.config['MAIL_SERVER'] = 'smtp.gmail.com'  # Use your SMTP server details
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_DEFAULT_SENDER'] = 'vedikapatil2713@gmail.com'
app.config['MAIL_USERNAME'] = 'vedikapatil2713@gmail.com'  # Your email address
app.config['MAIL_PASSWORD'] = 'iemz htjj jrxz thfc'  # Your email password
mail = Mail(app)


# Database configuration
DB_CONFIG = {
    "host": "localhost",
    "user": "root",
    "password": "password@456",
    "database": "question_bank",
}

# Function to create a database if it doesn't exist
def create_database():
    try:
        connection = mysql.connector.connect(
            host=DB_CONFIG["host"],
            user=DB_CONFIG["user"],
            password=DB_CONFIG["password"]
        )
        cursor = connection.cursor()
        cursor.execute(f"CREATE DATABASE IF NOT EXISTS {DB_CONFIG['question_bank']}")
        print(f"Database '{DB_CONFIG['question_bank']}' is ready.")
        cursor.close()
        connection.close()
    except Error as e:
        print(f"Error creating database: {e}")

# Function to establish a database connection
def create_connection():
    try:
        return mysql.connector.connect(**DB_CONFIG)
    except Error as e:
        print(f"Error connecting to MySQL: {e}")
        return None

# Function to initialize tables
def initialize_database():
    connection = create_connection()
    if connection:
        try:
            cursor = connection.cursor()

            # ✅ Create Users table
            cursor.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id INT AUTO_INCREMENT PRIMARY KEY,
                email VARCHAR(255) UNIQUE NOT NULL,
                username VARCHAR(255) UNIQUE NOT NULL,
                password VARCHAR(255) NOT NULL,
                role VARCHAR(50) NOT NULL,
                otp VARCHAR(50),
                status ENUM('pending', 'approved', 'rejected') DEFAULT 'pending',
                is_verified BOOLEAN DEFAULT FALSE
            );
            """)

            # ✅ Create Subjects table
            cursor.execute("""
            CREATE TABLE IF NOT EXISTS subjects (
                subject_id INT AUTO_INCREMENT PRIMARY KEY,
                subject_name VARCHAR(100) NOT NULL,
                branch VARCHAR(100) NOT NULL,
                semester INT NOT NULL,
                UNIQUE (subject_name, branch)
            );
            """)

            # ✅ Create Questions table
            cursor.execute("""
            CREATE TABLE IF NOT EXISTS questions (
                question_id INT AUTO_INCREMENT PRIMARY KEY,
                question_text TEXT NOT NULL,
                rbt_level INT NOT NULL,
                co INT NOT NULL,
                pi INT NOT NULL,
                marks INT NOT NULL,
                subject_id INT NOT NULL,
                user_id INT NOT NULL, 
                image_path VARCHAR(255) NULL,
                FOREIGN KEY (subject_id) REFERENCES subjects(subject_id) ON DELETE CASCADE,
                FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
            );
            """)

            print("Tables created successfully!")
        except Error as e:
            print(f"Error creating tables: {e}")
        finally:
            cursor.close()
            connection.close()

 # Routes
@app.route('/')
def index():
    return render_template('index.html')


# ✅ Function to generate a JWT token
def generate_verification_token(email):
    return jwt.encode(
        {'email': email, 'exp': datetime.datetime.utcnow() + datetime.timedelta(hours=1)}, 
        app.secret_key, algorithm="HS256"
    )

# ✅ Function to send a verification email
def send_verification_email(email):
    token = generate_verification_token(email)
    verification_url = url_for('verify_email', token=token, _external=True)

    msg = Message("Verify Your Email", recipients=[email])
    msg.body = f"Click the link to verify your account: {verification_url}"
    mail.send(msg)

def is_valid_password(password):
    """Check if the password meets the security criteria"""
    pattern = re.compile(r"^(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*?&])[A-Za-z\d@$!%*?&]{5,}$")
    return pattern.match(password)

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm-password']
        role = request.form['role']

        # Password Validation
        if password != confirm_password:
            flash("Passwords do not match!", "danger")
            return redirect(url_for('signup'))

        if not is_valid_password(password):
            flash("Password must be at least 8 characters long, include an uppercase letter, a number, and a special character.", "danger")
            return redirect(url_for('signup'))

        hashed_password = generate_password_hash(password)  # Hash password before storing

        connection = create_connection()
        cursor = connection.cursor()

        if role == "Admin":
            cursor.execute("INSERT INTO users (username, email, password, role, is_verified) VALUES (%s, %s, %s, %s, %s)", 
                           (username, email, hashed_password, role, False))
            connection.commit()

            # Send verification email (implement this function)
            send_verification_email(email)
            flash("A verification email has been sent. Please check your inbox.", "info")

        else:
            cursor.execute("INSERT INTO users (username, email, password, role, is_verified) VALUES (%s, %s, %s, %s, %s)", 
                           (username, email, hashed_password, role, True))
            connection.commit()
            flash("Signup successful! Please log in.", "success")

        cursor.close()
        connection.close()
        return redirect(url_for('login'))

    return render_template('signup.html')


@app.route('/admin_user')
def admin_user():
    connection = create_connection()
    cursor = connection.cursor()
    
    cursor.execute("SELECT id, username, email, role, status FROM users WHERE status = 'pending'")
    users = cursor.fetchall()
    
    cursor.close()
    connection.close()
    
    return render_template('admin_user.html', users=users)

@app.route('/update_user_status', methods=['POST'])
def update_user_status():
    user_id = request.form['user_id']
    action = request.form['action']

    new_status = "approved" if action == "approve" else "rejected"

    connection = create_connection()
    cursor = connection.cursor()
    
    cursor.execute("UPDATE users SET status = %s WHERE id = %s", (new_status, user_id))
    connection.commit()
    
    cursor.close()
    connection.close()

    flash(f"User status updated to {new_status}.", "success")
    return redirect(url_for('admin_user'))




@app.route('/verify-email/<token>')
def verify_email(token):
    try:
        data = jwt.decode(token, app.secret_key, algorithms=["HS256"])
        email = data['email']

        connection = create_connection()
        cursor = connection.cursor()
        cursor.execute("UPDATE users SET is_verified = TRUE WHERE email = %s", (email,))
        connection.commit()
        cursor.close()
        connection.close()

        flash("Email verified successfully! You can now log in.", "success")
    except jwt.ExpiredSignatureError:
        flash("Verification link has expired. Please sign up again.", "danger")
    except jwt.InvalidTokenError:
        flash("Invalid verification link.", "danger")

    return redirect(url_for('login'))

#login
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        connection = create_connection()
        cursor = connection.cursor(dictionary=True)  # ✅ Fetch as a dictionary

        cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()

    

        if user:
            stored_hash = user['password']  # Get hashed password from DB
            if check_password_hash(stored_hash, password):  # Compare hashed password
                if user['status'] == 'approved':  # Only allow login if approved
                    session['user_id'] = user['id']
                    flash("Login successful!", "success")

                    if user['role'] == 'Teacher':
                        return redirect(url_for('add_question'))
                    elif user['role'] == 'Higher Authority':
                        return redirect(url_for('generate_question_paper'))
                    elif user['role'] == 'Admin':
                        return redirect(url_for('admin_user'))
                elif user['status'] == 'pending':
                    flash("Your account is pending approval. Please wait for admin approval.", "warning")
                elif user['status'] == 'rejected':
                    flash("Your account has been rejected. You do not have access to login.", "danger")
                return redirect(url_for('login'))
            else:
                flash("Invalid username or password", "danger")
        else:
            flash("Invalid username or password", "danger")

            
        cursor.close()
        connection.close()
    return render_template('login.html')











def login_required(f):
    """Decorator to restrict access to logged-in users."""
    def wrapper(*args, **kwargs):
        if 'username' not in session:  # If no user is logged in
            return redirect(url_for('login'))  # Redirect to login page
        return f(*args, **kwargs)  # Otherwise, proceed to the page
    wrapper.__name__ = f.__name__
    return wrapper





















# Route to display the "Forgot Password" form
@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email')
        
        # Check if the email exists in the database
        connection = create_connection()
        if connection:
            try:
                cursor = connection.cursor()
                cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
                user = cursor.fetchone()

                if user:
                    # Generate OTP using pyotp
                    otp = f"{random.randint(1000, 9999)}" # Generate a random OTP

                    # Store the OTP in the database temporarily (you can store it for a limited time in your real application)
                    cursor.execute("UPDATE users SET otp = %s WHERE email = %s", (otp, email))
                    connection.commit()

                    # Send the OTP to the user's email
                    msg = Message("Your OTP for Password Reset", recipients=[email])
                    msg.body = f"Your OTP is: {otp}. Use this OTP to reset your password."
                    mail.send(msg)

                    flash('OTP sent to your email address.', 'info')
                    return redirect(url_for('verify_otp', email=email))  # Redirect to OTP verification page
                else:
                    flash('Email not found!', 'error')

            except Error as e:
                flash(f"Error: {e}", 'error')
            finally:
                cursor.close()
                connection.close()
        else:
            flash('Failed to connect to the database.', 'error')
        
    return render_template('forgot_password.html')


# Route to verify OTP and reset the password
@app.route('/verify_otp', methods=['GET', 'POST'])
def verify_otp():
    email = request.args.get('email')  # Get email from the URL
    if request.method == 'POST':
        otp = request.form.get('otp')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')

        if not otp or not new_password or not confirm_password:
            flash('All fields are required!', 'error')
            return redirect(url_for('verify_otp', email=email))

        if new_password != confirm_password:
            flash('Passwords do not match!', 'error')
            return redirect(url_for('verify_otp', email=email))

        # Validate OTP and reset the password
        connection = create_connection()
        if connection:
            try:
                cursor = connection.cursor()
                cursor.execute("SELECT otp FROM users WHERE email = %s", (email,))
                stored_otp = cursor.fetchone()

                if stored_otp and stored_otp[0] == otp:
                    # Hash the new password before storing it
                    hashed_password = generate_password_hash(new_password, method="pbkdf2:sha256", salt_length=8)

                    # Update the user's password in the database
                    cursor.execute("UPDATE users SET password = %s WHERE email = %s", (hashed_password, email))
                    connection.commit()
                    flash('Password reset successful!', 'success')
                    return redirect(url_for('login'))
                else:
                    flash('Invalid OTP!', 'error')

            except Error as e:
                flash(f"Error: {e}", 'error')
            finally:
                cursor.close()
                connection.close()
        else:
            flash('Failed to connect to the database.', 'error')

    return render_template('verify_otp.html', email=email)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
            
#add_question
# Route to render the HTML form and handle form submissions
# Route to render the HTML form and handle form submissions
@app.route('/add_question', methods=['GET', 'POST'])
def add_question():
    if 'user_id' not in session:
        flash("Please log in.", "warning")
        return redirect(url_for('login'))

    connection = create_connection()
    cursor = connection.cursor()

    # Fetch existing questions for the user
    cursor.execute("SELECT question_text FROM questions WHERE user_id = %s", (session['user_id'],))
    existing_questions = [row[0].strip().lower() for row in cursor.fetchall()]

    if request.method == "POST":
        question_text = request.form["question_text"].strip().lower()
        matched_questions = request.form.get("matched_questions", "[]")  # Default to "[]" if empty

        try:
            matched_questions = json.loads(matched_questions) if matched_questions else []
        except json.JSONDecodeError:
            matched_questions = []

        # If there are matched questions, show a flash message
        if matched_questions:
            flash("Question already exists in matched results!", "danger")
            return redirect(url_for('add_question'))

        # Proceed with adding the question if no match
        branch = request.form["branch"]
        semester = request.form["semester"]
        subject = request.form["subject"]
        rbt_level = request.form["rbt_level"]
        co = request.form["co"]
        pi = request.form["pi"]
        marks = request.form["marks"]
        user_id = session['user_id']

        # Save image if uploaded
        image_path = None
        if 'question_image' in request.files:
            image_file = request.files['question_image']
            if image_file.filename:
                filename = secure_filename(image_file.filename)
                image_path = os.path.join(UPLOAD_FOLDER, filename)
                image_file.save(image_path)

        # Get or insert subject_id
        cursor.execute("SELECT subject_id FROM subjects WHERE subject_name = %s AND branch = %s AND semester = %s",
                       (subject, branch, semester))
        subject_result = cursor.fetchone()

        if not subject_result:
            cursor.execute("INSERT INTO subjects (subject_name, branch, semester) VALUES (%s, %s, %s)",
                           (subject, branch, semester))
            connection.commit()
            subject_id = cursor.lastrowid
        else:
            subject_id = subject_result[0]

        # Insert new question
        cursor.execute("""
            INSERT INTO questions (question_text, rbt_level, co, pi, marks, subject_id, user_id, image_path) 
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """, (question_text, rbt_level, co, pi, marks, subject_id, user_id, image_path))
        connection.commit()

        flash("Question added successfully!", "success")

        return redirect(url_for('add_question'))

    cursor.close()
    connection.close()

    return render_template("add_question.html")

@app.route('/fetch_questions')
def fetch_questions():
    
    if 'user_id' not in session:
        return jsonify([])  # Return empty if not logged in

    user_id = session['user_id']
    connection = create_connection()
    cursor = connection.cursor(dictionary=True)

    cursor.execute("SELECT question_text FROM questions WHERE user_id = %s", (user_id,))
    questions = [row['question_text'] for row in cursor.fetchall()]
    
    cursor.close()
    connection.close()
    return jsonify(questions)



@app.route('/show_questions')
def show_questions():
    if 'user_id' not in session:
        flash("Please log in.", "warning")
        return redirect(url_for('login'))

    user_id = session['user_id']

    connection = create_connection()
    cursor = connection.cursor()
    
    cursor.execute("SELECT question_id, question_text, marks, rbt_level, co, pi, image_path FROM questions WHERE user_id = %s", (user_id,))
    questions = cursor.fetchall()  # Returns a list of tuples

    connection.close()
    
    return render_template('show_questions.html', questions=questions)



@app.route('/get_subjects', methods=['GET'])
def get_subjects():
    user_id = session.get('user_id')  # Get logged-in user's ID
    if not user_id:
        return jsonify({"error": "User not logged in"}), 401

    conn = create_connection()
    cursor = conn.cursor()

    # Step 1: Fetch subject IDs from questions where user_id matches
    cursor.execute("SELECT DISTINCT subject_id FROM questions WHERE user_id = ?", (user_id,))
    subject_ids = [row[0] for row in cursor.fetchall()]

    if not subject_ids:
        return jsonify([])  # No subjects found

    # Step 2: Fetch subject names based on subject IDs
    query = f"SELECT DISTINCT name FROM subjects WHERE id IN ({','.join(['?'] * len(subject_ids))})"
    cursor.execute(query, subject_ids)
    subjects = [row[0] for row in cursor.fetchall()]

    conn.close()
    return jsonify(subjects)  # Return only names



@app.route("/update_question", methods=["POST"])
def update_question():
    question_id = request.form.get("question_id")  # Ensure this matches the form input name
    question_text = request.form.get("question")  # Use 'question_text' as per schema
    marks = request.form.get("marks")
    rbt_level = request.form.get("rbt_level")
    co = request.form.get("co")
    pi = request.form.get("pi")
    
    

    
    if not question_id or not question_text:  
        flash("Missing required fields!", "danger")
        return redirect(url_for("show_questions"))

    conn = create_connection()
    cursor = conn.cursor()

    # Handle Image Upload (if a new image is uploaded)
    if "image" in request.files and request.files["image"].filename:
        image = request.files["image"]
        image_filename = f"question_{question_id}.jpg"
        image_path = os.path.join(app.config["UPLOAD_FOLDER"], image_filename)
        image.save(image_path)

        cursor.execute("""
            UPDATE questions 
            SET question_text = %s, marks = %s, rbt_level = %s, co = %s, pi = %s, image_path = %s 
            WHERE question_id = %s
        """, (question_text, marks, rbt_level, co, pi, image_filename, question_id))
    else:
        cursor.execute("""
            UPDATE questions 
            SET question_text = %s, marks = %s, rbt_level = %s, co = %s, pi = %s 
            WHERE question_id = %s
        """, (question_text, marks, rbt_level, co, pi, question_id))

    conn.commit()
    conn.close()

    flash("Question updated successfully!", "success")
    return redirect(url_for("show_questions"))




# Route to delete a question
@app.route("/delete_question/<int:question_id>")
def delete_question(question_id):
    conn = create_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM questions WHERE question_id = %s", (question_id,))
    conn.commit()
    conn.close()

    flash("Question deleted successfully.", "success")
    return redirect(url_for("show_questions"))



@app.route('/generate_question_paper', methods=['GET', 'POST'])
def generate_question_paper():
    if request.method == 'POST':
        # ✅ Extract form data safely
        branch = request.form.get('branch')
        semester = request.form.get('semester')
        subject = request.form.get('subject')
        subject_code = request.form.get('subject_code')
        exam_time = request.form.get('exam_time')
        exam_date = request.form.get('exam_date')
        examination = request.form.get('examination')
        paper_code = request.form.get('paper_code')

        # ✅ Ensure total_marks is an integer
        total_marks = request.form.get('total_marks', '0')
        try:
            total_marks = int(total_marks)
        except ValueError:
            flash("Invalid total marks entered.", 'error')
            return redirect(url_for('generate_question_paper'))

        # ✅ Database connection
        connection = create_connection()
        if not connection:
            flash("Database connection failed!", 'error')
            return redirect(url_for('generate_question_paper'))

        try:
            cursor = connection.cursor()

            # ✅ Fetch `subject_id` for the given subject
            cursor.execute("SELECT subject_id FROM subjects WHERE subject_name = %s", (subject,))
            subject_result = cursor.fetchone()

            if not subject_result:
                flash("Subject not found in database!", 'error')
                return redirect(url_for('generate_question_paper'))

            subject_id = subject_result[0]

            # ✅ Fetch questions for the subject (random order) with image_path
            cursor.execute("""
                SELECT question_text, CAST(marks AS SIGNED), rbt_level, co, pi, image_path 
                FROM questions 
                WHERE subject_id = %s
                ORDER BY RAND()
            """, (subject_id,))
            all_questions = cursor.fetchall()  

            if not all_questions:
                flash("No questions found for the selected subject.", 'error')
                return redirect(url_for('generate_question_paper'))

            cursor.close()
            connection.close()

            # ✅ Select questions that sum up to total_marks
            selected_questions = []
            current_marks = 0

            for question in all_questions:
                if current_marks + question[1] <= total_marks:
                    selected_questions.append(question)
                    current_marks += question[1]
                if current_marks == total_marks:
                    break

            if current_marks != total_marks:
                flash("Could not select exact marks, closest possible paper generated.", 'warning')

            # ✅ Generate Word Document
            doc = Document()

            # ✅ College Header
            header = doc.add_paragraph()
            header.alignment = 1
            header.add_run("D Y Patil College of Engineering, ").bold = True
            header.add_run("Salokhenagar, Kolhapur 416007\n").bold = True
            header.add_run("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING (DATA SCIENCE)\n").bold = True

            # ✅ Exam Details Table
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'

            cells = table.rows[0].cells
            cells[0].text = examination
            cells[1].text = f"Sem : {semester}"
            cells[2].text = f"Subject : {subject}"
            cells[3].text = f"Subject Code: {subject_code}"

            cells = table.rows[1].cells
            cells[0].text = f"Date : {exam_date}"
            cells[1].text = f"Maximum Marks: {total_marks}"
            cells[2].text = f"QP Code: {paper_code}"
            cells[3].text = "All questions are compulsory."

            # ✅ Add Separator Line
            line_paragraph = doc.add_paragraph()
            line_run = line_paragraph.add_run("_" * 95)
            line_run.bold = True
            line_paragraph.alignment = 1 
            # ✅ Questions Table (Without Separate Image Column)
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'

            # Table Headers
            headers = ["Q. No.", "Question", "Marks", "RBT Level", "CO", "PI", "Marks Secured"]
            for i, header_text in enumerate(headers):
                table.rows[0].cells[i].text = header_text

            # ✅ Add Selected Questions to Table (With Image Below Question)
            for i, (question_text, marks, rbt_level, co, pi, image_path) in enumerate(selected_questions, start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)  # Question No.
    
                # ✅ Add Question Text
                question_paragraph = row_cells[1].paragraphs[0]
                question_paragraph.add_run(question_text)

                # ✅ Add Image Below Question (if Exists)
                if image_path and os.path.exists(image_path):
                    question_paragraph.add_run("\n")  # Line break before image
                    run = question_paragraph.add_run("\n")
                    run.add_picture(image_path, width=Inches(2.0))  # Adjust width as needed

                row_cells[2].text = str(marks)
                row_cells[3].text = str(rbt_level)
                row_cells[4].text = str(co)
                row_cells[5].text = str(pi)
                row_cells[6].text = ""  # Marks Secured (empty for now)


                

            # ✅ Save Document in Memory and Send as Download
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return send_file(file_stream, as_attachment=True, download_name=f"{subject}.docx")

        except mysql.connector.Error as e:
            flash(f"Database error: {e}", 'error')

        finally:
            if connection.is_connected():
                connection.close()

    return render_template('generate_question_paper.html')







@app.route('/logout')
def logout():
    # Clear the session to log out the user
    session.clear()
    flash('You have been logged out successfully.', 'info')
    return redirect(url_for('login'))


 
# Initialize database and tables
initialize_database()

if __name__ == "__main__":
    app.run(debug=True)
    

